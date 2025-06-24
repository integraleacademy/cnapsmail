
from flask import Flask, render_template, request, redirect, url_for, send_file
import os
import sqlite3
import json
from datetime import datetime
from docx import Document

app = Flask(__name__)
app.secret_key = "cnaps_clé_ultra_secrète_42"

DB_NAME = "cnaps.db"
JSON_DATA_PATH = "data.json"
UPLOAD_FOLDER = "static/uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ---------- FORMULAIRE PUBLIC ----------
@app.route("/submit", methods=["POST"])
def submit():
    nom = request.form["nom"]
    prenom = request.form["prenom"]
    email = request.form["email"]
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    data = []
    if os.path.exists(JSON_DATA_PATH):
        with open(JSON_DATA_PATH, "r", encoding="utf-8") as f:
            try:
                data = json.load(f)
            except:
                pass

    data.append({"nom": nom, "prenom": prenom, "email": email, "timestamp": now})
    with open(JSON_DATA_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    return redirect("/?submitted=true")

@app.route("/")
def accueil():
    submitted = request.args.get("submitted", False)
    return render_template("formulaire.html", submitted=submitted)

# ---------- INTERFACE ADMIN ----------
@app.route("/admin")
def admin():
    filtre_cnaps = request.args.get("filtre_cnaps", "Tous")
    with sqlite3.connect(DB_NAME) as conn:
        conn.row_factory = sqlite3.Row
        cur_statuts = conn.execute("SELECT DISTINCT statut_cnaps FROM dossiers")
        statuts_disponibles = sorted([row["statut_cnaps"] for row in cur_statuts if row["statut_cnaps"]])
        if filtre_cnaps != "Tous":
            cur = conn.execute("SELECT * FROM dossiers WHERE statut_cnaps=?", (filtre_cnaps,))
        else:
            cur = conn.execute("SELECT * FROM dossiers")
        dossiers = cur.fetchall()
    return render_template("index.html", dossiers=dossiers, filtre_cnaps=filtre_cnaps, statuts_disponibles=statuts_disponibles)

@app.route("/statut/<int:id>", methods=["POST"])
def update_statut(id):
    statut = request.form.get("statut")
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("UPDATE dossiers SET statut = ? WHERE id = ?", (statut, id))
    return redirect("/admin")

@app.route("/statut_cnaps/<int:id>", methods=["POST"])
def update_statut_cnaps(id):
    statut_cnaps = request.form.get("statut_cnaps")
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("UPDATE dossiers SET statut_cnaps = ? WHERE id = ?", (statut_cnaps, id))
    return redirect("/admin")

@app.route("/commentaire/<int:id>", methods=["POST"])
def update_commentaire(id):
    commentaire = request.form.get("commentaire")
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("UPDATE dossiers SET commentaire = ? WHERE id = ?", (commentaire, id))
    return redirect("/admin")

@app.route("/supprimer/<int:id>")
def supprimer(id):
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("DELETE FROM dossiers WHERE id = ?", (id,))
    return redirect("/admin")

@app.route("/attestation/<int:id>")
def attestation(id):
    with sqlite3.connect(DB_NAME) as conn:
        conn.row_factory = sqlite3.Row
        cur = conn.execute("SELECT * FROM dossiers WHERE id = ?", (id,))
        dossier = cur.fetchone()

    doc = Document()
    doc.add_heading("Attestation de Préinscription", 0)
    doc.add_paragraph(f"{dossier['prenom']} {dossier['nom']} est préinscrit(e) à la formation {dossier['formation']} le {dossier['session']}.")
    filename = f"attestation_{dossier['nom']}_{dossier['prenom']}.docx"
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    doc.save(filepath)
    return send_file(filepath, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=10000)
